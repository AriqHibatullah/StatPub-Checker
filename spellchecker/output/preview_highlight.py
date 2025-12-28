from __future__ import annotations

import io, re
from typing import Iterable, List, Tuple
from copy import deepcopy

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX

def _iter_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def highlight_terms_docx_bytes(
    docx_bytes: bytes,
    terms: List[str],
    *,
    case_insensitive: bool = True,
    whole_word: bool = True,
    color: WD_COLOR_INDEX = WD_COLOR_INDEX.YELLOW,
    min_len: int = 3,
) -> bytes:
    toks = sorted({(t or "").strip() for t in terms if (t or "").strip()}, key=len, reverse=True)
    toks = [t for t in toks if len(t) >= min_len]
    if not toks:
        return docx_bytes

    flags = re.IGNORECASE if case_insensitive else 0
    if whole_word:
        pat = re.compile(r"\b(" + "|".join(re.escape(t) for t in toks) + r")\b", flags)
    else:
        pat = re.compile("(" + "|".join(re.escape(t) for t in toks) + ")", flags)

    doc = Document(io.BytesIO(docx_bytes))

    for p in _iter_paragraphs(doc):
        runs = p.runs
        if not runs:
            continue

        full = "".join(r.text or "" for r in runs)
        if not full:
            continue

        matches = list(pat.finditer(full))
        if not matches:
            continue

        hl_ranges: List[Tuple[int, int]] = [(m.start(), m.end()) for m in matches]

        base_rpr = deepcopy(runs[0]._r.get_or_add_rPr())

        for r in list(p.runs):
            r._r.getparent().remove(r._r)

        # merge overlap
        merged = []
        for a, b in sorted(hl_ranges):
            if not merged or a > merged[-1][1]:
                merged.append([a, b])
            else:
                merged[-1][1] = max(merged[-1][1], b)
        hl_ranges = [(a, b) for a, b in merged]

        cursor = 0
        for a, b in hl_ranges:
            if a > cursor:
                r = p.add_run(full[cursor:a])
                r._r.get_or_add_rPr().clear()
                r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

            r = p.add_run(full[a:b])
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))
            r.font.highlight_color = color

            cursor = b

        if cursor < len(full):
            r = p.add_run(full[cursor:])
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
