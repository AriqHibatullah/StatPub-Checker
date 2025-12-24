from __future__ import annotations
import io, re, fitz
from typing import Dict, Iterable, Tuple, List
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
import pandas as pd

def transfer_case(src: str, dst: str) -> str:
    if not src or not dst:
        return dst
    if src.isupper():
        return dst.upper()
    if len(src) >= 2 and src[0].isupper() and src[1:].islower():
        return dst[:1].upper() + dst[1:]
    return dst

def _iter_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def replace_and_highlight_docx_bytes(
    docx_bytes: bytes,
    replacements: Dict[str, str],
    *,
    case_insensitive: bool = False,
    whole_word: bool = True,
    color: WD_COLOR_INDEX = WD_COLOR_INDEX.YELLOW,
) -> bytes:
    if not replacements:
        return docx_bytes

    items = [(k, v) for k, v in replacements.items() if k and v and k != v]
    items.sort(key=lambda kv: len(kv[0]), reverse=True)
    if not items:
        return docx_bytes

    flags = re.IGNORECASE if case_insensitive else 0

    olds = [re.escape(k) for k, _ in items]
    if whole_word:
        pat = re.compile(r"\b(" + "|".join(olds) + r")\b", flags)
    else:
        pat = re.compile("(" + "|".join(olds) + r")", flags)

    if case_insensitive:
        repl_map = {k.lower(): v for k, v in items}
        def repl_func(m: re.Match) -> str:
            found = m.group(0)
            new = repl_map.get(found.lower())
            if not new:
                return found
            return transfer_case(found, new)
    else:
        repl_map = dict(items)
        def repl_func(m: re.Match) -> str:
            return repl_map.get(m.group(0), m.group(0))

    doc = Document(io.BytesIO(docx_bytes))

    for p in _iter_paragraphs(doc):
        runs = p.runs
        if not runs:
            continue

        full = "".join(r.text or "" for r in runs)
        if not full:
            continue

        matches = list(pat.finditer(full))
        if not matches:
            continue

        spans: List[Tuple[int, int, int]] = []
        pos = 0
        for i, r in enumerate(runs):
            t = r.text or ""
            spans.append((i, pos, pos + len(t)))
            pos += len(t)

        new_full_parts = []
        hl_ranges: List[Tuple[int, int]] = []
        last = 0
        out_pos = 0

        for m in matches:
            if m.start() > last:
                seg = full[last:m.start()]
                new_full_parts.append(seg)
                out_pos += len(seg)

            old_txt = m.group(0)
            new_txt = repl_func(m)

            new_full_parts.append(new_txt)
            hl_ranges.append((out_pos, out_pos + len(new_txt)))
            out_pos += len(new_txt)

            last = m.end()

        if last < len(full):
            seg = full[last:]
            new_full_parts.append(seg)

        new_full = "".join(new_full_parts)

        if new_full == full:
            continue

        base_rpr = deepcopy(runs[0]._r.get_or_add_rPr())

        for r in list(p.runs):
            r._r.getparent().remove(r._r)

        merged = []
        for a, b in sorted(hl_ranges):
            if not merged or a > merged[-1][1]:
                merged.append([a, b])
            else:
                merged[-1][1] = max(merged[-1][1], b)
        hl_ranges = [(a, b) for a, b in merged]

        cursor = 0
        for a, b in hl_ranges:
            if a > cursor:
                txt = new_full[cursor:a]
                r = p.add_run(txt)
                r._r.get_or_add_rPr().clear()
                r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

            txt = new_full[a:b]
            r = p.add_run(txt)
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))
            r.font.highlight_color = color

            cursor = b

        if cursor < len(new_full):
            txt = new_full[cursor:]
            r = p.add_run(txt)
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def highlight_terms_docx_bytes(
    docx_bytes: bytes,
    terms: List[str],
    *,
    case_insensitive: bool = True,
    whole_word: bool = True,
    color: WD_COLOR_INDEX = WD_COLOR_INDEX.YELLOW,
    min_len: int = 3,
) -> bytes:
    toks = sorted({(t or "").strip() for t in terms if (t or "").strip()}, key=len, reverse=True)
    toks = [t for t in toks if len(t) >= min_len]
    if not toks:
        return docx_bytes

    flags = re.IGNORECASE if case_insensitive else 0
    if whole_word:
        pat = re.compile(r"\b(" + "|".join(re.escape(t) for t in toks) + r")\b", flags)
    else:
        pat = re.compile("(" + "|".join(re.escape(t) for t in toks) + ")", flags)

    doc = Document(io.BytesIO(docx_bytes))

    for p in _iter_paragraphs(doc):
        runs = p.runs
        if not runs:
            continue

        full = "".join(r.text or "" for r in runs)
        if not full:
            continue

        matches = list(pat.finditer(full))
        if not matches:
            continue

        hl_ranges: List[Tuple[int, int]] = [(m.start(), m.end()) for m in matches]

        base_rpr = deepcopy(runs[0]._r.get_or_add_rPr())

        for r in list(p.runs):
            r._r.getparent().remove(r._r)

        # merge overlap
        merged = []
        for a, b in sorted(hl_ranges):
            if not merged or a > merged[-1][1]:
                merged.append([a, b])
            else:
                merged[-1][1] = max(merged[-1][1], b)
        hl_ranges = [(a, b) for a, b in merged]

        cursor = 0
        for a, b in hl_ranges:
            if a > cursor:
                r = p.add_run(full[cursor:a])
                r._r.get_or_add_rPr().clear()
                r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

            r = p.add_run(full[a:b])
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))
            r.font.highlight_color = color

            cursor = b

        if cursor < len(full):
            r = p.add_run(full[cursor:])
            r._r.get_or_add_rPr().clear()
            r._r.get_or_add_rPr().extend(deepcopy(base_rpr))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
    
def locate_tokens_in_pdf_pages(pdf_bytes: bytes, tokens: list[str], *, min_len: int = 3):
    toks = []
    for t in tokens:
        t = (t or "").strip()
        if len(t) >= min_len:
            toks.append(t)
    toks = sorted(set(toks))
    if not toks:
        return pd.DataFrame(columns=["kata", "page", "hits"])

    tokset = {t.lower() for t in toks}

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    rows = []

    for i in range(len(doc)):
        page = doc[i]
        page_num = i + 1
        words = page.get_text("words")

        hits = {}
        for w in words:
            text = w[4] or ""
            norm = re.sub(r"^[^\w'’\-]+|[^\w'’\-]+$", "", text).lower()
            if norm in tokset:
                hits[norm] = hits.get(norm, 0) + 1

        for kw, n in hits.items():
            rows.append({"kata": kw, "page": page_num, "hits": n})

    df = pd.DataFrame(rows, columns=["kata", "page", "hits"])
    return df


