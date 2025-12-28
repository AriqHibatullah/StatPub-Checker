from __future__ import annotations

import io
import re
import uuid
import zipfile
import tempfile
from docx import Document
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

import pandas as pd
import streamlit as st

from spellchecker.vocab.loaders import load_kbbi_words, load_txt_set
from spellchecker.vocab.load_storage import load_resources_from_storage
from spellchecker.pipeline import run_on_file, build_vocabs
from spellchecker.settings import Settings

from spellchecker.extractors.docx_extractor import docx_bytes_to_pdf_bytes
from spellchecker.session.ensure import ensure_session_state, sync_uploaded_files_and_autoreset
from spellchecker.session.review_helpers import apply_maps_to_df, commit_from_editor_state

from spellchecker.output.docx_highlighter import transfer_case, replace_and_highlight_docx_bytes, highlight_terms_docx_bytes, locate_tokens_in_pdf_pages
from spellchecker.output.notifier_resend import send_dev_report_email
from spellchecker.output.reporter import SupabaseConfig, upload_dev_run_report

# =========================
# Streamlit config
# =========================
st.set_page_config(page_title="StatPub Checker", layout="wide")
st.markdown("""
    <style>
    header[data-testid="stHeader"] {
        background-color: rgba(0,0,0,0) !important;
    }
    .stApp {
        background: linear-gradient(135deg, #f7fff7 0%, #c3cfe2 100%);
    }
    .stApp, .stApp * {
        color: #1a535c;
    }
    .stButton > button:not([disabled]) {
        background-color: #ffe66d !important;
        color: white !important;
        border: none !important;
    }
    .stButton > button:not([disabled]):hover {
        background-color: #ffe246 !important;
        color: white !important;
    }
    .stButton > button[kind="secondary"]:not([disabled]) {
        background-color: #f7fff7 !important;
        color: #1a535c !important;
        border: 1px solid rgba(26, 83, 92, 0.4) !important;
    }
    .stButton > button[kind="secondary"]:not([disabled]):hover {
        background-color: #F0F2F6 !important;
        color: #1a535c !important;
        border: 1px solid rgba(26, 83, 92, 0.4) !important;
    }
    [data-testid="stSidebar"] {
        background: #fafcfa !important;
    }
    section[data-testid="stSidebar"] {
        width: 300px !important;
    }
    section[data-testid="stSidebar"] > div {
        width: 300px !important;
    }
    </style>
""", unsafe_allow_html=True)

with st.sidebar:
    st.title("🔮 StatPub Checker 📃")
    st.caption("Lihat demo Web App StatPub Checker [di sini](https://docs.streamlit.io).")
    with st.expander("📘 Cara penggunaan"):
        st.markdown("""
            Panduan lengkap bagaimana cara menggunakan StatPub Checker tersedia [di sini](https://drive.google.com/file/d/1fFY97-FEgeOVvuf18r_1Ckvz6hHMjwD2/view?usp=sharing).
        """)
    with st.expander("ℹ️ Release Note Terbaru"):
        st.markdown("""
            ## v0.4.1\n
            - Menambahkan input Tipe publikasi untuk pengguna
            - Menghapus input Threshold Confidence untuk pengguna
            - Menambahkan sedikit penyesuaian pada program untuk developer
        """)
    st.info("Version 0.4.1")

    st.write(" ")
    st.write(" ")
    st.markdown("""
        <div style="padding: 10px; border-radius: 10px">
            <p style="margin:0; font-size:0.8rem; color:#1a535c;">Developed By:</p>
            <p style="margin:0; font-weight:bold; color:#4ecdc4;">Firdaini Azmi &  </p>
            <p style="margin:0; font-weight:bold; color:#4ecdc4;">Muhammad Ariq Hibatullah</p>
        </div>
    """, unsafe_allow_html=True)

st.title("📂 Masukkan file untuk diperiksa")
st.caption("Upload DOCX/PDF → sistem menghasilkan temuan typo + saran koreksi.")

# =========================
# Resource load
# =========================
ensure_session_state()
data_storage = st.secrets["STORAGE"]
ver = load_storage_version(data_storage)
resources = load_resources_from_storage_versioned(bucket=data_storage, version=ver)
EDITOR_KEY = "tabel_seleksi"

# =========================
# UI Controls
# =========================
uploads = st.file_uploader(
        "Upload file DOCX/PDF, bisa upload banyak",
        type=["docx", "pdf"],
        accept_multiple_files=True,
        key="uploader_files",
        help = "Untuk file pdf, disarankan agar file tidak mempunyai watermark untuk performa optimal."
    )
sync_uploaded_files_and_autoreset(uploads)

colA1, colA2, = st.columns([2,1])

with colA1:
    colB, colC = st.columns([1, 1])

    with colB:
        tipe_publikasi = st.selectbox(
            "Bahasa publikasi",
            ["Bahasa Indonesia", "Campuran"],
            help="Pilih tipe yang sesuai untuk memudahkan program memeriksa dokumen"
        )   
    
        show_only_top1_if_conf_ge = st.slider(
            "Skor Confidence untuk top-1",
            min_value=0.0,
            max_value=1.0,
            value=0.72,
            step=0.01,
            format="%.2f",
            help="Masukkan skor Confidence dari saran kata yang dianggap 'pasti benar'"
        )

    with colC:
        max_findings = st.number_input("Max temuan per file", min_value=50, max_value=5000, value=200, step=50)

with colA2:
    user_vocab_text = st.text_area(
            "Masukkan kata tambahan",
            height = 150,
            placeholder = "Contoh:\nStunting\nBig Data\n...",
            help = "Masukan kata yang khusus dan tidak ada di KBBI, seperti nama orang atau tempat"
        )

run_btn = st.button("Jalankan pemeriksaan", type="primary", disabled=not uploads)

if "df" not in st.session_state:
    st.session_state.df = None
if "report_ready" not in st.session_state:
    st.session_state.report_ready = False
if "csv_ready" not in st.session_state:
    st.session_state.csv_ready = False

def parse_user_vocab(text: str) -> set[str]:
    vocab = set()
    for line in (text or "").splitlines():
        w = line.strip().lower()
        if not w:
            continue
        for tok in w.split():
            if tok.isalpha() and len(tok) >= 2:
                vocab.add(tok)
    return vocab

user_vocab = parse_user_vocab(user_vocab_text)
st.session_state.user_vocab = sorted(list(user_vocab))

def findings_to_dataframe(findings: List[Any]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for f in findings:
        if isinstance(f, dict):
            file = f.get("file", "")
            page = f.get("page", "")
            token = f.get("token", "")
            status = f.get("status", "")
            snippet = f.get("snippet", "")
            suggs = f.get("suggestions", []) or []
        else:
            file = getattr(f, "file", "")
            page = getattr(f, "page", "")
            token = getattr(f, "token", "")
            status = getattr(f, "status", "")
            snippet = getattr(f, "snippet", "")
            suggs = getattr(f, "suggestions", []) or []

        row = {"file": file, "page": page, "token": token, "status": status}
        for i in range(1, 4):
            if len(suggs) >= i:
                s = suggs[i - 1]
                if isinstance(s, dict):
                    sug = s.get("suggestion") or s.get("term") or s.get("word") or ""
                    conf = s.get("confidence", "")
                else:
                    sug = getattr(s, "suggestion", "") or getattr(s, "term", "") or ""
                    conf = getattr(s, "confidence", "")
            else:
                sug, conf = "", ""
            row[f"suggestion_{i}"] = sug
            row[f"confidence_{i}"] = conf
        row["snippet"] = snippet
        rows.append(row)

    cols = [
        "file", "page", "token", "status",
        "suggestion_1", "confidence_1",
        "suggestion_2", "confidence_2",
        "suggestion_3", "confidence_3",
        "snippet",
    ]
    df = pd.DataFrame(rows)
    for c in cols:
        if c not in df.columns:
            df[c] = ""
    return df[cols]

def parse_id_ranges(text: str) -> set[int]:
    ids: set[int] = set()
    if not text:
        return ids

    parts = [p.strip() for p in text.split(",") if p.strip()]
    for part in parts:
        if "-" in part:
            a, b = [x.strip() for x in part.split("-", 1)]
            if a.isdigit() and b.isdigit():
                start, end = int(a), int(b)
                if start > end:
                    start, end = end, start
                ids.update(range(start, end + 1))
        else:
            if part.isdigit():
                ids.add(int(part))
    return ids

def run_pipeline_on_paths(paths: List[str]) -> List[Any]:
    cfg = Settings(
        topk=int(3),
        max_findings_per_file=int(max_findings),
        show_only_top1_if_conf_ge=float(show_only_top1_if_conf_ge),
    )

    known_vocab_plus = set(resources["known_vocab"]) | set(user_vocab or set())

    all_findings: List[Any] = []
    for p in paths:
        findings, meta = run_on_file(
            path=p,
            cfg=cfg,
            known_vocab=known_vocab_plus,
            english_vocab=resources["english_vocab"],
            known_vocab_for_names=resources["known_vocab_for_names"],
            ignore_vocab=resources["ignore_vocab"],
            domain_terms=resources["domain_terms"],
            protected_phrases=resources["protected_phrases"],
            protected_name_tokens=resources["protected_name_tokens"],
        )
        all_findings.extend(findings)

    return all_findings

def replace_in_docx_bytes(docx_bytes: bytes, replacements: dict[str, str]) -> bytes:
    items = [(k, v) for k, v in replacements.items() if k and v and k != v]
    if not items:
        return docx_bytes

    items.sort(key=lambda kv: len(kv[0]), reverse=True)

    pattern = re.compile(r"\b(" + "|".join(re.escape(k) for k, _ in items) + r")\b", re.IGNORECASE)
    repl_lower = {k.lower(): v for k, v in items}

    def sub_func(m: re.Match) -> str:
        found = m.group(0)
        new = repl_lower.get(found.lower())
        if not new:
            return found
        return transfer_case(found, new)

    doc = Document(io.BytesIO(docx_bytes))

    def repl_paragraph(p) -> None:
        full = "".join(r.text or "" for r in p.runs)
        if not full:
            return

        new_full = pattern.sub(sub_func, full)
        if new_full == full:
            return

        for r in list(p.runs):
            r._r.getparent().remove(r._r)

        p.add_run(new_full)

    for p in doc.paragraphs:
        repl_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl_paragraph(p)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

if run_btn:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        saved_paths: List[str] = []

        if "upload_bytes_by_name" not in st.session_state:
            st.session_state.upload_bytes_by_name = {}

        for up in uploads:
            out_path = tmpdir / up.name
            b = up.getbuffer().tobytes()
            out_path.write_bytes(b)
            saved_paths.append(str(out_path))

            st.session_state.upload_bytes_by_name[up.name] = b

        st.info(f"Memproses {len(saved_paths)} file…")

        with st.spinner("Running spellcheck…"):
            findings = run_pipeline_on_paths(saved_paths)

        df = findings_to_dataframe(findings)

        st.session_state.run_id = str(uuid.uuid4())
        st.session_state.run_ts_utc = datetime.utcnow().isoformat()

        df_raw_dev = df.copy()
        if "_rid" not in df_raw_dev.columns:
            df_raw_dev["_rid"] = range(len(df_raw_dev))
        st.session_state.df_raw_dev = df_raw_dev

        if "_rid" not in df.columns:
            df["_rid"] = range(len(df))

        st.session_state.df = df
        st.session_state.report_ready = True
        st.session_state.review_mode = False
        st.session_state.csv_ready = False

def _do_preview():
    st.session_state.preview_show = True
    st.session_state.preview_file = file_pilih

if st.session_state.report_ready and st.session_state.df is not None:
    df_raw = st.session_state.df

    STATUS_MAP = {
        "symspell": "🔴 Kesalahan penulisan",
        "no_candidates": "⚪ Kata tidak dikenali",
        "space_error": "🟡 Kesalahan spasi",
        "capital_error": "🟡 Kapital awal kalimat",
        "abbr_candidate": "⚪ Singkatan tidak dikenali",
        "confusion": "🔴 Kesalahan penulisan",
        "affix_typo": "🔴 Kesalahan penulisan",
    }
    df_view = df_raw[df_raw["status"] != "abbr_confirmed"].copy()
    if "status" in df_view.columns:
        df_view["status"] = df_view["status"].map(STATUS_MAP).fillna(df_view["status"])

    st.subheader("Report Pemeriksaan")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total temuan", int(len(df_view)))
    c2.metric("Jumlah file", int(df_view["file"].nunique()) if "file" in df_view.columns else 0)
    
    TOTAL_TYPO_LABEL = {"🔴 Kesalahan penulisan", "🟡 Kesalahan spasi", "🟡 Kapital awal kalimat"}
    total_typo = int(df_view["status"].isin(TOTAL_TYPO_LABEL).sum()) if "status" in df_view.columns else 0
    c3.metric("Total kesalahan penulisan", total_typo)

    UNKNOWN_LABELS = {"⚪ Kata tidak dikenali", "⚪ Singkatan tidak dikenali"}
    total_unknown = int(df_view["status"].isin(UNKNOWN_LABELS).sum()) if "status" in df_view.columns else 0
    c4.metric("Total tidak dikenali", total_unknown)

    tab3, tab1, tab2 = st.tabs(["Preview Detail", "Ringkasan", "Per File"])

    with tab1:
        if "status" in df_view.columns and len(df_view) > 0:
            st.markdown("**Ringkasan status**")
            status_summary = (
                df_view["status"]
                .value_counts(dropna=False)
                .rename_axis("Status")
                .reset_index(name="Jumlah")
            )
            st.dataframe(status_summary, width='stretch')

        st.markdown("**Kata yang paling sering muncul**")
        if "token" in df_view.columns and len(df_view) > 0:
            top_tokens = (
                df_view["token"]
                .value_counts(dropna=False)
                .head(10)
                .rename_axis("Kata")
                .reset_index(name="Jumlah")
            )
            st.dataframe(top_tokens, width='stretch')

    with tab2:
        st.markdown("**Lihat preview dokumen**")
        file_pilih = None
        if "pdf_cache_by_name" not in st.session_state:
            st.session_state.pdf_cache_by_name = {}
            
        if "file" in df_view.columns:
            options = sorted(df_view["file"].dropna().astype(str).unique())
            options = ["— Pilih file —"] + options
            file_pilih = st.selectbox("Pilih file", options) if len(options) else None
            
            if file_pilih == "— Pilih file —":
                file_pilih = None

        preview_btn = st.button("Tampilkan preview", type="secondary", on_click=_do_preview)
        if st.session_state.preview_show and st.session_state.preview_file:
            file_to_render = st.session_state.preview_file
            b = st.session_state.upload_bytes_by_name.get(file_to_render)

            if b is None:
                st.error("File tidak ditemukan. Silakan upload ulang atau jalankan proses lagi.")
            else:
                colsB, colsC = st.columns([2, 1])

                df_file = df_view[df_view["file"].astype(str) == str(file_pilih)]
                tokens_file = df_file["token"].dropna().astype(str).tolist()
        
                with colsB:
                    with st.spinner("Menyiapkan preview dokumen..."):
                        if file_pilih.lower().endswith(".pdf"):
                            st.pdf(b, height=600)

                            with colsC:
                                st.markdown("**Kata temuan**")
                                df_pages = locate_tokens_in_pdf_pages(b, tokens_file)
                                if df_pages.empty:
                                    st.info("Tidak ada match ditemukan (mungkin PDF scan / tidak ada text layer).")
                                else:
                                    df_pages = df_pages.sort_values(["kata", "page"])
                                    st.dataframe(df_pages, width='stretch', height=600)
                                    
                        elif file_pilih.lower().endswith(".docx"):
                            try:
                                docx_hl = highlight_terms_docx_bytes(
                                    b,
                                    tokens_file,
                                    case_insensitive=True,
                                    whole_word=True,
                                )
                            except Exception as e:
                                st.error(f"Gagal highlight DOCX: {e}")
                                docx_hl = b

                            cached_pdf = st.session_state.pdf_cache_by_name.get(file_pilih)
                            if cached_pdf is None:
                                try:
                                    pdf_bytes = docx_bytes_to_pdf_bytes(docx_hl)
                                    st.session_state.pdf_cache_by_name[file_pilih] = pdf_bytes
                                    cached_pdf = pdf_bytes
                                except Exception as e:
                                    st.error(f"Gagal convert DOCX ke PDF: {e}")
                                    pdf_bytes = None
                                    cached_pdf = None
        
                            if cached_pdf is not None:
                                st.pdf(cached_pdf, height=600)
        
                                with colsC:
                                    st.markdown("**Kata temuan**")
                                    df_pages = locate_tokens_in_pdf_pages(cached_pdf, tokens_file)
                                    if df_pages.empty:
                                        st.info("Tidak ada match ditemukan (PDF hasil konversi tidak punya text layer yang terdeteksi).")
                                    else:
                                        df_pages = (
                                            df_pages.sort_values("page", ascending=True)[["kata", "page"]]
                                            .reset_index(drop=True)
                                        )
                                        st.dataframe(df_pages, width='stretch', height=500)
        
                        else:
                            st.error("Format file tidak didukung.")

        st.markdown("**Temuan per file**")
        if "file" in df_view.columns and len(df_view) > 0:
            by_file = (
                df_view.groupby("file", dropna=False)
                .size()
                .sort_values(ascending=False)
                .rename("Jumlah")
                .reset_index()
            )
            st.dataframe(by_file, width='stretch')

    with tab3:
        st.markdown("**Preview detail**")
        st.caption("Silahkan diperiksa terlebih dahulu hasil deteksi di bawah, terlebih lagi yang berstatus ⚪ (tidak dikenali).")
        DISPLAY_COLS = [
            "file",
            "token",
            "status",
            "suggestion_1",
            "suggestion_2",
            "suggestion_3",
            "snippet",
        ]
        pre_detail = (
            df_view
            .loc[:, DISPLAY_COLS]
            .rename(columns={
                "file": "File",
                "token": "Kata",
                "suggestion_1": "Saran 1",
                "suggestion_2": "Saran 2",
                "suggestion_3": "Saran 3",
                "status": "Status",
                "snippet": "Pada kalimat"
            })
        )
        st.dataframe(pre_detail.head(500), width='stretch', height=520)

        if "token" in df_view.columns and len(df_view) > 0:
            vc = df_view["token"].value_counts(dropna=False)
            vc = vc[vc > 5]

            if not vc.empty:
                st.write(" ")
                st.markdown("**Kata ini sering muncul, apakah merupakan sebuah kata yang benar?**")
                st.caption("Jika memang merupakan kata yang benar, centanglah kata tersebut saat seleksi pilihan kata.")

                sering_muncul = (
                    vc.head(10)
                    .rename_axis("Kata")
                    .reset_index(name="Jumlah")
                )
                st.dataframe(sering_muncul, width='stretch')

    st.markdown("---")
    start_review = st.button("Mulai Seleksi Pilihan Kata", type="primary")
    if start_review:
        st.session_state.review_mode = True

if st.session_state.get("review_mode", False) and st.session_state.df is not None:
    df_raw = st.session_state.df.copy()

    status_map = {
        "symspell": "🔴 Kesalahan penulisan",
        "no_candidates": "⚪ Kata tidak dikenali",
        "space_error": "🟡 Kesalahan spasi",
        "capital_error": "🟡 Kapital awal kalimat",
        "abbr_candidate": "⚪ Singkatan tidak dikenali",
        "confusion": "🔴 Kesalahan penulisan",
        "affix_typo": "🔴 Kesalahan penulisan",
    }
    df = df_raw[df_raw["status"] != "abbr_confirmed"].copy()

    if "fix_choice" not in df.columns:
        df["fix_choice"] = ""
    if "fix_custom" not in df.columns:
        df["fix_custom"] = ""
    if "fix_final" not in df.columns:
        df["fix_final"] = ""
    if "ignore" not in df.columns:
        df["ignore"] = False

    st.subheader("Review & Seleksi")
    st.write("Silahkan seleksi jika ada kata yang salah koreksi oleh sistem. Centanglah kata yang bukan typo/salah koreksi.")
    st.caption(
        "Segmen pada table: 🟩 Fix; 🟨 Review.  \n"  
        "Bantu perkembangan sistem ini dengan mengisi kolom pada segmen 🟨 Review, agar sistem dapat lebih mengoreksi lebih optimal."
    )

    cols = st.columns(3)
    with cols[0]:
        file_filter = (
            st.multiselect("Filter file", sorted(df["file"].dropna().unique().tolist()))
            if "file" in df.columns else []
        )
    with cols[1]:
        if "status" in df.columns:
            df["status_disp"] = df["status"].map(status_map).fillna(df["status"].astype(str))
        else:
            df["status_disp"] = ""

        status_filter = (
            st.multiselect("Filter status", sorted(df["status_disp"].dropna().unique().tolist()))
            if "status_disp" in df.columns else []
        )
    with cols[2]:
        st.write(" ")
        st.write(" ")
        show_only_fixable = st.checkbox("Hanya tampilkan yang akan diperbaiki", value=False)

    review_cols = ["ignore", "fix_choice", "fix_custom"]
    info_cols = [c for c in ["status_disp", "token", "suggestion_1", "suggestion_2", "suggestion_3"] if c in df.columns]

    filtered = df
    if file_filter and "file" in filtered.columns:
        filtered = filtered[filtered["file"].isin(file_filter)]
    if status_filter and "status_disp" in filtered.columns:
        filtered = filtered[filtered["status_disp"].isin(status_filter)]
    if show_only_fixable:
        filtered = filtered[filtered["ignore"] == False]

    display_cols = ["_rid"] + info_cols + review_cols
    view = filtered[display_cols].copy()
    view = apply_maps_to_df(view)

    st.session_state["_rid_order_for_editor"] = view["_rid"].astype(int).tolist()

    edited = st.data_editor(
        view,
        key=EDITOR_KEY,
        on_change=commit_from_editor_state,
        hide_index=True,
        width='stretch',
        height=520,
        num_rows="fixed",
        column_config={
            "token": st.column_config.TextColumn("Kata"),
            "status_disp": st.column_config.TextColumn("Status kesalahan"),
            "suggestion_1": st.column_config.TextColumn("🟦 Saran 1"),
            "suggestion_2": st.column_config.TextColumn("🟦 Saran 2"),
            "suggestion_3": st.column_config.TextColumn("🟦 Saran 3"),
            "ignore": st.column_config.CheckboxColumn(
                "🟩 Salah koreksi",
                help="Centang jika ini BUKAN typo dan tidak perlu diperbaiki atau koreksi salah",
                default=False,
            ),
            "fix_choice": st.column_config.SelectboxColumn(
                "🟩 Pilih koreksi",
                help="Pilih Saran 1-3 (jika tersedia) atau '➕ Manual'.",
                options=["Saran 1", "Saran 2", "Saran 3", "➕ Manual"],
            ),
            "fix_custom": st.column_config.TextColumn(
                "🟩 Koreksi manual",
                help="Isi jika pilih 'custom'",
            ),
            "_rid": st.column_config.NumberColumn("ID", disabled=True),
        },
        disabled=[c for c in view.columns if c not in review_cols]
    )
    st.session_state.salah_koreksi.update(dict(zip(edited["_rid"].astype(int), edited["ignore"].astype(bool))))
    st.session_state.pilihan_koreksi.update(dict(zip(edited["_rid"].astype(int), edited["fix_choice"].astype(str))))
    st.session_state.koreksi_manual.update(dict(zip(edited["_rid"].astype(int), edited["fix_custom"].astype(str))))

    base = df.set_index("_rid")
    ed = edited.set_index("_rid")

    base.loc[ed.index, review_cols] = ed[review_cols]
    df = base.reset_index()

    AUTO_IGNORE_STATUSES = {"no_candidates", "abbr_candidate"}

    def _final_fix(row):    
        if row.get("status") in AUTO_IGNORE_STATUSES:
            return ""

        if bool(row.get("ignore", False)):
            return ""

        choice = str(row.get("fix_choice", "") or "").strip()
        custom = str(row.get("fix_custom", "") or "").strip()

        if choice == "":
            if custom:
                return custom
            return str(row.get("suggestion_1", "") or "").strip()

        if choice in ("1", "2", "3"):
            return str(row.get(f"suggestion_{choice}", "") or "").strip()

        if choice in ("suggestion_1", "suggestion_2", "suggestion_3"):
            return str(row.get(choice, "") or "").strip()

        if custom:
            return custom
        return str(row.get("suggestion_1", "") or "").strip()
    
    df["fix_final"] = df.apply(_final_fix, axis=1)
    st.session_state.df = df

    st.caption("Setelah seleksi selesai, isi review di bawah dan klik Perbaiki typo untuk membuat output berdasarkan pilihanmu.")
    
    reason_options = ["", "Kata yg benar", "Saran salah", "Nama/istilah khusus", "Singkatan", "Bahasa campuran", "Lainnya"]
    selected_rids = [rid for rid, v in st.session_state.salah_koreksi.items() if v]

    with st.popover("🟨 Review salah koreksi"):
        if not selected_rids:
            st.info("Tidak ada kata yang salah koreksi.")
        else:
            df_selected = view[
                view["_rid"].isin(selected_rids)
            ].copy().sort_values("_rid")

            for _, row in df_selected.iterrows():
                rid = int(row["_rid"])
                kata = row["token"]

                if rid not in st.session_state.review_alasan:
                    st.session_state.review_alasan[rid] = "-"
                if rid not in st.session_state.review_catatan:
                    st.session_state.review_catatan[rid] = ""

                st.markdown(f"**Kata:** {kata}")

                c1, c2 = st.columns([1, 1])
                with c1:
                    alasan = st.selectbox(
                        "Alasan",
                        options=reason_options,
                        index=reason_options.index(st.session_state.review_alasan[rid])
                        if st.session_state.review_alasan[rid] in reason_options else 0,
                        key=f"alasan_{rid}",
                        label_visibility="collapsed",
                    )
                with c2:
                    catatan = st.text_input(
                        "Catatan",
                        value=st.session_state.review_catatan[rid],
                        key=f"catatan_{rid}",
                        label_visibility="collapsed",
                        placeholder="Catatan (opsional)...",
                    )

                st.session_state.review_alasan[rid] = alasan
                st.session_state.review_catatan[rid] = catatan

                st.divider()

            st.caption(f"Total dicentang: {len(selected_rids)}")
    
    st.markdown("---")
    
    highlight_out = st.checkbox("Sertakan highlight pada draft hasil", value=True, key="highlight_out")
    fix_btn = st.button("Perbaiki typo", type="secondary")

    if fix_btn:
        with st.spinner("Membuat dokumen revisi + log perubahan..."):
            df_all = st.session_state.df.copy()

            run_id = st.session_state.get("run_id") or str(uuid.uuid4())
            st.session_state.run_id = run_id
            date_str = datetime.utcnow().strftime("%Y-%m-%d")
            ts_utc = st.session_state.get("run_ts_utc") or datetime.utcnow().isoformat()

            df_all["action"] = df_all.apply(
                lambda r: "ignored" if r.get("ignore", False)
                else (
                    "replaced"
                    if (str(r.get("fix_final", "")).strip()
                        and str(r.get("fix_final", "")).strip() != str(r.get("token", "")).strip())
                    else "no_fix"
                ),
                axis=1
            )

            df_audit_user = df_all[df_all["action"] == "replaced"].copy()
            cols_user = [c for c in ["file", "token", "fix_final", "snippet"] if c in df_audit_user.columns]
            df_audit_user = df_audit_user.loc[:, cols_user].rename(columns={
                "file": "File",
                "token": "Kata sebelum",
                "fix_final": "Kata sesudah",
                "snippet": "Pada Kalimat",
            })
            user_log_bytes = df_audit_user.to_csv(index=False).encode("utf-8")

            df_fix = df_all[
                (df_all["ignore"] == False) &
                (df_all["fix_final"].astype(str).str.strip().str.len() > 0)
            ].copy()

            if "status" in df_all.columns:
                df_fix = df_fix[~df_fix["status"].isin({"no_candidates", "abbr_candidate"})]

            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as z:
                z.writestr("log_perubahan.csv", user_log_bytes)

                groups = list(df_fix.groupby("file"))
                prog = st.progress(0)
                total = max(len(groups), 1)

                for i, (fname, g) in enumerate(groups, start=1):
                    prog.progress(i / total)

                    src_bytes = st.session_state.upload_bytes_by_name.get(fname)
                    if not src_bytes:
                        continue

                    repl = {}
                    for _, r in g.iterrows():
                        old = str(r.get("token", "")).strip()
                        new = str(r.get("fix_final", "")).strip()
                        if old and new and old != new:
                            repl[old] = new

                    if fname.lower().endswith(".docx"):
                        if highlight_out:
                            revised = replace_and_highlight_docx_bytes(
                                src_bytes,
                                repl,
                                case_insensitive=True,
                                whole_word=True,
                            )
                        else:
                            revised = replace_in_docx_bytes(src_bytes, repl)

                        z.writestr(f"revisi_{fname}", revised)
                    else:
                        z.writestr(
                            f"README_{fname}.txt",
                            "File ini bukan DOCX, revisi otomatis belum didukung. Lihat log_perubahan.csv"
                        )

                prog.empty()

            zip_buf.seek(0)

            st.download_button(
                "Download hasil revisi (ZIP)",
                data=zip_buf.getvalue(),
                file_name="hasil_revisi.zip",
                mime="application/zip",
            )

        try:
            df_raw_dev = st.session_state.get("df_raw_dev")
            if df_raw_dev is None:
                df_raw_dev = df_all.copy()

            df_eval_full = df_all.copy()

            meta = {
                "run_id": run_id,
                "ts_utc": ts_utc,
                "files": sorted(list(st.session_state.get("upload_bytes_by_name", {}).keys())),
                "config": {
                    "topk": int(3),
                    "max_findings": int(max_findings),
                    "show_only_top1_if_conf_ge": float(show_only_top1_if_conf_ge),
                },
                "summary": {
                    "total_findings": int(len(df_eval_full)),
                    "replaced": int((df_eval_full["action"] == "replaced").sum()) if "action" in df_eval_full.columns else None,
                    "ignored": int((df_eval_full["action"] == "ignored").sum()) if "action" in df_eval_full.columns else None,
                },
                "user_vocab": st.session_state.get("user_vocab", []),
                "user_vocab_count": len(st.session_state.get("user_vocab", [])),
                "app_version": "0.4.1",
            }

            bucket = "dev-reports"
            base_path = f"{date_str}/{run_id}"

            cfg_sb = SupabaseConfig(
                url=st.secrets["URL"],
                service_role_key=st.secrets["ROLE_KEY"],
            )

            uv = "\n".join(st.session_state.get("user_vocab", [])) + "\n"

            upload_dev_run_report(
                cfg=cfg_sb,
                bucket=bucket,
                base_path=base_path,
                raw_findings_csv=df_raw_dev.to_csv(index=False).encode("utf-8"),
                eval_full_csv=df_eval_full.to_csv(index=False).encode("utf-8"),
                meta=meta,
                user_vocab_txt=uv.encode("utf-8"),
            )

            try:
                send_dev_report_email(
                    secrets=st.secrets,
                    run_id=run_id,
                    base_path=base_path,
                    total_findings=len(df_eval_full),
                )
            except Exception as e:
                st.warning(f"(Dev) Email notif gagal: {e}")

        except Exception as e:
            st.warning(f"(Dev) Upload report gagal: {e}")

st.markdown("---")
st.caption(
    "Catatan: Produk ini memakai kamus & model yang masih dikembangkan. "
    "Data akan selalu diupdate untuk memaksimalkan performa."
)
